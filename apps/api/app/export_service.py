# ruff: noqa: E501  (este módulo contém templates CSS/HTML com linhas longas por natureza)
from __future__ import annotations

import re
from dataclasses import dataclass
from html import escape
from io import BytesIO
from re import sub

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor

from app.text_repair import repair_text


@dataclass(frozen=True)
class ExportDocument:
    title: str
    subtitle: str
    metadata: list[tuple[str, str]]
    content: str
    filename_prefix: str


@dataclass(frozen=True)
class ExportResult:
    content: bytes
    media_type: str
    filename: str


EXPORT_FORMATS = {"pdf", "docx", "md", "html"}


def safe_filename(value: str) -> str:
    cleaned = sub(r"[^a-zA-Z0-9._-]+", "-", value.strip().lower()).strip("-")
    return cleaned[:90] or "duofy-document"


def export_document(document: ExportDocument, export_format: str) -> ExportResult:
    normalized_format = export_format.lower().strip()
    if normalized_format not in EXPORT_FORMATS:
        raise ValueError("Formato de exportacao nao suportado.")

    document = ExportDocument(
        title=repair_text(document.title),
        subtitle=repair_text(document.subtitle),
        metadata=[(repair_text(label), repair_text(value)) for label, value in document.metadata],
        content=repair_text(document.content),
        filename_prefix=document.filename_prefix,
    )
    filename = f"{safe_filename(document.filename_prefix)}.{normalized_format}"
    if normalized_format == "pdf":
        return ExportResult(
            content=build_duofy_pdf(document),
            media_type="application/pdf",
            filename=filename,
        )
    if normalized_format == "docx":
        return ExportResult(
            content=build_duofy_docx(document),
            media_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            filename=filename,
        )
    if normalized_format == "html":
        return ExportResult(
            content=build_duofy_html(document).encode("utf-8"),
            media_type="text/html; charset=utf-8",
            filename=filename,
        )
    return ExportResult(
        content=build_duofy_markdown(document).encode("utf-8"),
        media_type="text/markdown; charset=utf-8",
        filename=filename,
    )


def build_duofy_markdown(document: ExportDocument) -> str:
    metadata = "\n".join(f"- **{label}:** {value}" for label, value in document.metadata)
    blocks = [f"# {document.title}", document.subtitle]
    if metadata:
        blocks.extend(["## Metadados", metadata])
    blocks.append(document.content.strip())
    return "\n\n".join(block for block in blocks if block).strip() + "\n"


# ---------------------------------------------------------------------------
# Markdown -> HTML (compartilhado por HTML e PDF)
# ---------------------------------------------------------------------------


def markdown_to_html(content: str) -> str:
    html: list[str] = []
    ul_items: list[str] = []
    ol_items: list[str] = []
    table_rows: list[list[str]] = []

    def flush_ul() -> None:
        if ul_items:
            html.append("<ul>" + "".join(ul_items) + "</ul>")
            ul_items.clear()

    def flush_ol() -> None:
        if ol_items:
            html.append("<ol>" + "".join(ol_items) + "</ol>")
            ol_items.clear()

    def flush_table() -> None:
        if not table_rows:
            return
        header, *rows = table_rows
        head = "".join(f"<th>{_inline_markdown(cell)}</th>" for cell in header)
        body = "".join(
            "<tr>" + "".join(f"<td>{_inline_markdown(cell)}</td>" for cell in row) + "</tr>"
            for row in rows
        )
        html.append(
            f'<table class="md"><thead><tr>{head}</tr></thead><tbody>{body}</tbody></table>'
        )
        table_rows.clear()

    def flush_all() -> None:
        flush_ul()
        flush_ol()
        flush_table()

    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            flush_all()
            continue
        if line in {"---", "***", "___"}:
            flush_all()
            html.append("<hr>")
            continue
        if _is_markdown_table_row(line):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                continue
            flush_ul()
            flush_ol()
            table_rows.append(cells)
            continue
        ordered = re.match(r"^\d+[.)]\s+(.*)$", line)
        if ordered:
            flush_ul()
            flush_table()
            ol_items.append(f"<li>{_inline_markdown(ordered.group(1).strip())}</li>")
            continue
        if line.startswith(("- ", "* ", "+ ")):
            flush_ol()
            flush_table()
            ul_items.append(f"<li>{_inline_markdown(line[2:].strip())}</li>")
            continue
        flush_all()
        if line.startswith("#### ") or line.startswith("##### "):
            html.append(f"<h4>{_inline_markdown(line.lstrip('#').strip())}</h4>")
        elif line.startswith("### "):
            html.append(f"<h3>{_inline_markdown(line[4:])}</h3>")
        elif line.startswith("## "):
            html.append(f"<h2>{_inline_markdown(line[3:])}</h2>")
        elif line.startswith("# "):
            html.append(f"<h2>{_inline_markdown(line[2:])}</h2>")
        elif line.startswith(">"):
            html.append(f"<blockquote>{_inline_markdown(line.lstrip('>').strip())}</blockquote>")
        else:
            html.append(f"<p>{_inline_markdown(line)}</p>")
    flush_all()
    return "\n".join(html)


def _inline_markdown(value: str) -> str:
    escaped = escape(value)
    # links [texto](url) — só http(s)/mailto viram âncora
    def _link(match: re.Match[str]) -> str:
        text, url = match.group(1), match.group(2)
        if re.match(r"^(https?:|mailto:)", url, re.IGNORECASE):
            return f'<a href="{url}">{text}</a>'
        return text

    escaped = re.sub(r"\[([^\]]+)\]\(([^)\s]+)\)", _link, escaped)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"__(.+?)__", r"<strong>\1</strong>", escaped)
    escaped = re.sub(r"(?<![\w*])\*(?!\s)(.+?)(?<!\s)\*(?![\w*])", r"<em>\1</em>", escaped)
    escaped = re.sub(r"(?<![\w_])_(?!\s)(.+?)(?<!\s)_(?![\w_])", r"<em>\1</em>", escaped)
    escaped = re.sub(r"`(.+?)`", r"<code>\1</code>", escaped)
    return escaped


def _is_markdown_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|") and line.count("|") >= 2


# ---------------------------------------------------------------------------
# HTML de tela (export .html)
# ---------------------------------------------------------------------------

_SCREEN_STYLE = """
    body { margin: 0; background: #f7f7fb; color: #11131a;
      font-family: ui-sans-serif, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; }
    main { width: min(880px, calc(100% - 48px)); margin: 48px auto; background: #fff;
      border: 1px solid #e9e8ef; border-radius: 24px;
      box-shadow: 0 24px 70px rgba(18,20,30,.08); padding: 56px; }
    .brand { color: #6d35ee; font-weight: 800; letter-spacing: -.04em; }
    h1 { font-size: 40px; line-height: 1.06; letter-spacing: -.04em; margin: 12px 0 8px; }
    h2 { font-size: 21px; letter-spacing: -.02em; margin-top: 32px; border-bottom: 1px solid #eceaf4; padding-bottom: 6px; }
    h3 { font-size: 16px; margin-top: 22px; }
    p, li, td, th { font-size: 15px; line-height: 1.7; }
    .subtitle { color: #6b7280; margin-bottom: 26px; }
    table { width: 100%; border-collapse: collapse; margin: 18px 0 28px; table-layout: fixed; }
    th, td { border: 1px solid #e9e8ef; padding: 10px 12px; text-align: left; vertical-align: top; overflow-wrap: anywhere; }
    thead th { background: #f6f3ff; color: #5b29d6; }
    .meta th { width: 200px; background: #f4efff; }
    blockquote { margin: 16px 0; padding: 12px 18px; border-left: 4px solid #6d35ee; background: #f8f6ff; border-radius: 0 8px 8px 0; }
    a { color: #6d35ee; }
    code { font-family: "SFMono-Regular", Consolas, monospace; background: #f4efff; color: #5b29d6; padding: 1px 5px; border-radius: 4px; }
"""


def build_duofy_html(document: ExportDocument) -> str:
    metadata_rows = "\n".join(
        f"<tr><th>{escape(label)}</th><td>{escape(value)}</td></tr>"
        for label, value in document.metadata
    )
    meta_block = (
        f'<table class="meta">{metadata_rows}</table>' if metadata_rows else ""
    )
    return f"""<!doctype html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{escape(document.title)}</title>
  <style>{_SCREEN_STYLE}</style>
</head>
<body>
  <main>
    <div class="brand">Duofy</div>
    <h1>{escape(document.title)}</h1>
    <p class="subtitle">{escape(document.subtitle)}</p>
    {meta_block}
    {markdown_to_html(document.content)}
  </main>
</body>
</html>
"""


# ---------------------------------------------------------------------------
# PDF profissional via WeasyPrint (HTML/CSS -> PDF, paginado, sem cortes)
# ---------------------------------------------------------------------------

_PDF_STYLE = """
  @page {
    size: A4;
    margin: 20mm 16mm 18mm 16mm;
    @top-left { content: "DUOFY"; font: 700 8pt "Liberation Sans", sans-serif; color: #6d35ee; letter-spacing: 1.5px; }
    @top-right { content: string(doctitle); font: 8pt "Liberation Sans", sans-serif; color: #aab0bd; }
    @bottom-left { content: "Documento gerado pela plataforma Duofy"; font: 7.5pt "Liberation Sans", sans-serif; color: #bcc0cb; }
    @bottom-right { content: "Página " counter(page) " / " counter(pages); font: 7.5pt "Liberation Sans", sans-serif; color: #aab0bd; }
  }
  @page :first {
    @top-left { content: ""; }
    @top-right { content: ""; }
  }

  * { box-sizing: border-box; }
  html { font-family: "Liberation Sans", "DejaVu Sans", sans-serif; color: #1a1c25; font-size: 10.5pt; line-height: 1.62; }
  body { margin: 0; }

  .cover { border-bottom: 3px solid #6d35ee; padding-bottom: 14px; margin-bottom: 22px; }
  .wordmark { font-weight: 800; font-size: 11pt; letter-spacing: 0.5px; color: #6d35ee; text-transform: uppercase; }
  h1.title { font-size: 24pt; line-height: 1.12; letter-spacing: -0.02em; margin: 9px 0 6px; color: #11131a;
    string-set: doctitle content(text); }
  .subtitle { color: #6b7280; font-size: 11pt; margin: 0; line-height: 1.45; }

  .meta { width: 100%; border-collapse: collapse; table-layout: fixed; margin: 0 0 24px;
    border: 1px solid #e9e8ef; border-radius: 8px; overflow: hidden; }
  .meta th { width: 34%; background: #f6f3ff; color: #5b29d6; font-weight: 700; text-align: left;
    padding: 8px 12px; font-size: 8.6pt; vertical-align: top; border-bottom: 1px solid #efedf6; overflow-wrap: anywhere; }
  .meta td { padding: 8px 12px; font-size: 9.2pt; color: #1a1c25; vertical-align: top;
    border-bottom: 1px solid #efedf6; overflow-wrap: anywhere; }
  .meta tr:last-child th, .meta tr:last-child td { border-bottom: 0; }

  h2 { font-size: 14pt; letter-spacing: -0.01em; margin: 22px 0 8px; padding-bottom: 5px;
    border-bottom: 1px solid #eceaf4; color: #11131a; break-after: avoid; }
  h3 { font-size: 11.5pt; margin: 16px 0 5px; color: #2a2140; break-after: avoid; }
  h4 { font-size: 10pt; margin: 12px 0 4px; color: #4a4560; text-transform: uppercase; letter-spacing: 0.4px; break-after: avoid; }
  p { margin: 0 0 8px; orphans: 2; widows: 2; }
  strong { color: #11131a; font-weight: 700; }
  em { font-style: italic; }
  a { color: #6d35ee; text-decoration: none; }
  code { font-family: "DejaVu Sans Mono", monospace; background: #f4efff; color: #5b29d6;
    padding: 1px 4px; border-radius: 3px; font-size: 8.6pt; }

  ul, ol { margin: 6px 0 12px; padding: 0; list-style: none; }
  li { margin: 0 0 5px; padding-left: 18px; position: relative; orphans: 2; widows: 2; }
  ul > li::before { content: ""; position: absolute; left: 3px; top: 6.5px; width: 5px; height: 5px;
    border-radius: 50%; background: #8b5cf6; }
  ol { counter-reset: item; }
  ol > li { counter-increment: item; }
  ol > li::before { content: counter(item) "."; position: absolute; left: 0; top: 0;
    color: #6d35ee; font-weight: 700; font-size: 9.5pt; }

  blockquote { margin: 12px 0; padding: 9px 16px; border-left: 3px solid #6d35ee;
    background: #f8f6ff; border-radius: 0 6px 6px 0; color: #3a3550; break-inside: avoid; }

  table.md { width: 100%; border-collapse: collapse; table-layout: fixed; margin: 12px 0 18px; font-size: 8.8pt; }
  table.md thead { display: table-header-group; }
  table.md th { background: #f6f3ff; color: #5b29d6; text-align: left; font-weight: 700; }
  table.md th, table.md td { border: 1px solid #e9e8ef; padding: 6px 8px; vertical-align: top;
    overflow-wrap: anywhere; word-break: break-word; line-height: 1.4; }
  table.md tr { break-inside: avoid; }

  hr { border: 0; border-top: 1px solid #e9e8ef; margin: 18px 0; }
"""


def _pdf_html(document: ExportDocument) -> str:
    metadata_rows = "\n".join(
        f"<tr><th>{escape(label)}</th><td>{escape(value)}</td></tr>"
        for label, value in document.metadata
    )
    meta_block = f'<table class="meta">{metadata_rows}</table>' if metadata_rows else ""
    return f"""<!doctype html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <title>{escape(document.title)}</title>
  <style>{_PDF_STYLE}</style>
</head>
<body>
  <header class="cover">
    <div class="wordmark">Duofy</div>
    <h1 class="title">{escape(document.title)}</h1>
    <p class="subtitle">{escape(document.subtitle)}</p>
  </header>
  {meta_block}
  <section class="content">
    {markdown_to_html(document.content)}
  </section>
</body>
</html>
"""


def build_duofy_pdf(document_data: ExportDocument) -> bytes:
    # Import tardio: WeasyPrint depende de libs nativas presentes só no container da API.
    from weasyprint import HTML

    return HTML(string=_pdf_html(document_data)).write_pdf()


# ---------------------------------------------------------------------------
# DOCX (python-docx) — inalterado
# ---------------------------------------------------------------------------


def build_duofy_docx(document: ExportDocument) -> bytes:
    docx = DocxDocument()
    styles = docx.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(25)

    brand = docx.add_paragraph()
    brand_run = brand.add_run("Duofy")
    brand_run.bold = True
    brand_run.font.color.rgb = RGBColor(109, 53, 238)
    brand_run.font.size = Pt(13)

    docx.add_heading(document.title, level=0)
    subtitle = docx.add_paragraph(document.subtitle)
    subtitle.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    if document.metadata:
        table = docx.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in document.metadata:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = value
        docx.add_paragraph()

    pending_table: list[list[str]] = []

    def flush_docx_table() -> None:
        if not pending_table:
            return
        table = docx.add_table(rows=0, cols=len(pending_table[0]))
        table.style = "Table Grid"
        for cells in pending_table:
            row = table.add_row()
            for index, cell in enumerate(cells):
                row.cells[index].text = cell.replace("**", "")
        docx.add_paragraph()
        pending_table.clear()

    for raw_line in document.content.splitlines():
        line = raw_line.strip()
        if not line:
            flush_docx_table()
            docx.add_paragraph()
        elif line in {"---", "***"}:
            flush_docx_table()
            docx.add_paragraph()
        elif _is_markdown_table_row(line):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                continue
            pending_table.append(cells)
        elif line.startswith("### "):
            flush_docx_table()
            docx.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            flush_docx_table()
            docx.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            flush_docx_table()
            docx.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")):
            flush_docx_table()
            docx.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            flush_docx_table()
            docx.add_paragraph(line)
    flush_docx_table()

    buffer = BytesIO()
    docx.save(buffer)
    return buffer.getvalue()
